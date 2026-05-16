"""Generates 10 realistic mock HR policy documents (5 PDF, 5 DOCX)."""
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
import docx

OUTPUT_DIR = Path(__file__).parent


def make_pdf(filename: str, title: str, sections: list[tuple[str, str]]) -> None:
    path = OUTPUT_DIR / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 0.5*cm))
    for heading, body in sections:
        story.append(Paragraph(heading, styles["Heading2"]))
        story.append(Paragraph(body, styles["BodyText"]))
        story.append(Spacer(1, 0.3*cm))
    doc.build(story)
    print(f"Created {path}")


def make_docx(filename: str, title: str, sections: list[tuple[str, str]]) -> None:
    path = OUTPUT_DIR / filename
    document = docx.Document()
    document.add_heading(title, 0)
    for heading, body in sections:
        document.add_heading(heading, level=2)
        document.add_paragraph(body)
    document.save(str(path))
    print(f"Created {path}")


DOCS = [
    ("pdf", "pto_vacation_policy.pdf", "PTO & Vacation Policy", [
        ("Overview", "This policy outlines the paid time off (PTO) entitlements for all full-time employees of the company. Part-time employees receive PTO on a pro-rated basis."),
        ("Annual PTO Accrual", "Full-time employees accrue 15 days of PTO per calendar year, earned at a rate of 1.25 days per month. Employees may begin using PTO after completing their 90-day probationary period."),
        ("Sick Leave", "Employees are entitled to 10 days of sick leave per year, separate from PTO. Sick leave does not accrue and does not carry over to the following year. A medical certificate is required for absences exceeding 3 consecutive days."),
        ("Public Holidays", "The company observes 12 public holidays per year. Employees required to work on a public holiday will receive a day in lieu within 30 days."),
        ("PTO Rollover", "Unused PTO up to a maximum of 5 days may be carried over to the following calendar year. Any PTO in excess of 5 days will be forfeited on December 31st."),
        ("PTO Request Process", "Employees must submit PTO requests via the HR portal at least 2 weeks in advance for requests of 3 or more consecutive days. Requests of 1-2 days require 48 hours notice. Approval is subject to business needs and manager discretion."),
        ("PTO Payout on Termination", "Upon termination of employment, employees will be paid out for any accrued but unused PTO. Employees who resign without providing adequate notice may forfeit accrued PTO at the company's discretion."),
    ]),
    ("docx", "expense_reimbursement.docx", "Expense Reimbursement Policy", [
        ("Purpose", "This policy establishes guidelines for reimbursement of business expenses incurred by employees in the course of their duties. The company will reimburse reasonable and necessary expenses that are properly documented and approved."),
        ("Submission Deadline", "All expense claims must be submitted within 30 calendar days of the expense being incurred. Claims submitted after 30 days may be declined. Expense reports must be submitted through the company expense management system."),
        ("Meals and Entertainment", "The per diem meal allowance is EUR 50 per day for domestic travel and EUR 75 per day for international travel. Alcohol is not reimbursable unless part of a pre-approved client entertainment event. Business purpose must be documented for all meal expenses."),
        ("Accommodation", "Hotel expenses are reimbursable up to EUR 150 per night for domestic travel and EUR 200 per night for international travel. Employees should book through the company's preferred travel portal where possible. Any accommodation above these limits requires pre-approval from a director."),
        ("Receipt Requirements", "Original receipts are required for all expenses of EUR 25 or more. Credit card statements are not acceptable substitutes for itemized receipts. Receipts must clearly show the vendor, date, and amount."),
        ("Approval Thresholds", "Expenses up to EUR 500 require manager approval. Expenses between EUR 500 and EUR 2000 require director approval. Expenses above EUR 2000 require VP-level approval. Pre-approval is strongly recommended for large expenses."),
        ("Non-Reimbursable Expenses", "The following are not reimbursable: personal entertainment, fines and penalties, personal travel upgrades, gifts to employees, political donations, and expenses that violate company policy or applicable law."),
    ]),
    ("pdf", "remote_work_policy.pdf", "Remote Work Policy", [
        ("Eligibility", "Remote work arrangements are available to full-time employees who have completed their 6-month probationary period and whose role can be performed effectively outside the office. Eligibility is determined by the employee's manager in consultation with HR."),
        ("Remote Work Days", "Eligible employees may work remotely up to 3 days per week. The specific days must be agreed in advance with the employee's manager to ensure adequate team coverage and collaboration. At least 2 days per week must be worked from the office."),
        ("Core Hours", "All employees, whether working remotely or in the office, must be available and responsive during core hours of 10:00 AM to 3:00 PM in their local time zone. Meetings should be scheduled during these hours where possible."),
        ("Equipment and Technology", "The company provides a laptop for all employees. Employees working remotely are responsible for maintaining a reliable internet connection. The company will provide a monthly internet stipend of EUR 30 to help offset home internet costs."),
        ("Security Requirements", "Employees must use the company VPN whenever accessing internal company systems, databases, or confidential information remotely. Company devices must be kept physically secure. Employees must not use public Wi-Fi without the VPN."),
        ("Performance Expectations", "Remote work is a privilege, not a right. If an employee's performance declines or the arrangement creates business problems, the manager may require the employee to return to full-time office work. Remote work arrangements are reviewed annually."),
        ("Health and Safety", "Employees are responsible for ensuring their remote workspace is safe and ergonomically appropriate. The company may require employees to complete a home workspace assessment. The company is not liable for injuries that occur in a non-compliant home workspace."),
    ]),
    ("docx", "code_of_conduct.docx", "Code of Conduct", [
        ("Introduction", "This Code of Conduct sets out the standards of behavior expected from all employees, contractors, and third parties acting on behalf of the company. Adherence to this code is a condition of employment and business relationship."),
        ("Respectful Workplace", "All employees have the right to work in an environment free from harassment, discrimination, and bullying. We are committed to treating everyone with dignity and respect regardless of age, gender, race, ethnicity, religion, sexual orientation, disability, or any other protected characteristic."),
        ("Zero Tolerance for Harassment", "Sexual harassment, bullying, intimidation, and any form of workplace harassment are strictly prohibited. This applies in the workplace, at company events, and in work-related communications including email and messaging platforms. Violations will result in disciplinary action up to and including termination."),
        ("Conflicts of Interest", "Employees must avoid situations where personal interests conflict or appear to conflict with company interests. All potential conflicts of interest must be disclosed in writing to HR and your manager. This includes outside employment, financial interests in competitors or suppliers, and personal relationships with vendors."),
        ("Confidentiality", "Employees have access to confidential information about the company, clients, and colleagues. This information must not be shared outside the company without authorization. Confidentiality obligations continue after employment ends. Violation of confidentiality may result in legal action."),
        ("Use of Company Resources", "Company resources including technology, equipment, and time are for business purposes. Limited personal use of technology is permitted provided it does not interfere with work, violate this policy, or create security risks. Employees have no expectation of privacy on company systems."),
        ("Reporting Violations", "Employees who observe or suspect violations of this code are encouraged to report them to HR, their manager, or via the anonymous ethics hotline. The company prohibits retaliation against anyone who reports concerns in good faith. All reports will be investigated promptly and confidentially."),
    ]),
    ("pdf", "health_benefits.pdf", "Health and Benefits Policy", [
        ("Medical Insurance", "The company provides comprehensive medical insurance to all full-time employees and their eligible dependents. The company covers 80% of the premium cost; employees contribute the remaining 20% through payroll deduction. Coverage begins on the first day of the month following the employee's start date."),
        ("Dental Coverage", "Employees receive a dental allowance of EUR 500 per calendar year. This covers routine check-ups, fillings, and basic dental procedures. Orthodontic treatment is covered at 50% up to a lifetime maximum of EUR 2000. Claims are submitted directly to the dental insurance provider."),
        ("Vision Coverage", "The vision allowance is EUR 300 per calendar year, covering eye examinations, prescription glasses, and contact lenses. Laser eye surgery is not covered under the standard plan but may be covered under optional supplemental insurance."),
        ("Enrollment", "New employees must enroll in benefits within 30 days of their start date. Outside of new hire enrollment, changes can only be made during the annual open enrollment period in November, or following a qualifying life event such as marriage, birth of a child, or loss of other coverage."),
        ("Mental Health Support", "The company provides access to an Employee Assistance Program (EAP) which includes up to 8 free counseling sessions per year. The EAP also provides support for financial planning, legal advice, and work-life balance resources. The service is confidential and available 24/7."),
        ("Life and Disability Insurance", "All employees are covered by group life insurance equal to 2x their annual salary at no cost to the employee. Short-term disability insurance covers 70% of salary for up to 13 weeks. Long-term disability insurance covers 60% of salary after 13 weeks for up to age 65."),
        ("Wellness Program", "The company offers a wellness reimbursement of EUR 400 per year for gym memberships, fitness classes, or other qualified wellness activities. Employees may also participate in the company's subsidized healthy lunch program and access on-site meditation rooms."),
    ]),
    ("docx", "performance_review.docx", "Performance Review Process", [
        ("Overview", "The performance review process is designed to provide employees with regular, structured feedback on their performance, development, and alignment with company values. Reviews support decisions about compensation, promotions, and development planning."),
        ("Review Cycle", "Annual performance reviews are conducted in December and cover the full calendar year. A mid-year check-in is conducted in June to provide informal feedback and adjust goals as needed. The mid-year check-in does not result in a formal rating."),
        ("Rating Scale", "Performance is rated on a 5-point scale: 1 - Needs Improvement: performance is below expectations in key areas; 2 - Developing: approaching expectations, development needed; 3 - Meets Expectations: consistently meets role requirements; 4 - Exceeds Expectations: regularly goes beyond requirements; 5 - Exceptional: outstanding contribution significantly above expectations."),
        ("Self-Assessment", "All employees are required to complete a self-assessment before the annual review. The self-assessment covers achievements, challenges, goal progress, and development needs. Self-assessments must be submitted to managers at least one week before the review meeting."),
        ("Manager Assessment", "Managers complete a written assessment covering the same competencies as the self-assessment. For employees with 3 or more direct reports, a 360-degree feedback component is included, gathering input from peers and cross-functional stakeholders."),
        ("Promotions and Compensation", "Promotions are tied to sustained performance of 4 or above for at least 2 consecutive review cycles, demonstration of competencies at the next level, and role availability. Merit increases are determined based on the annual review rating and the company's compensation budget. Employees rated 1 are not eligible for merit increases."),
        ("Performance Improvement Plans", "Employees receiving a rating of 1 will be placed on a Performance Improvement Plan (PIP). The PIP outlines specific, measurable improvement targets and a timeline of 60-90 days. Failure to meet PIP targets may result in termination of employment."),
    ]),
    ("pdf", "onboarding_guide.pdf", "Employee Onboarding Guide", [
        ("Welcome", "Welcome to the company. This guide will help you navigate your first days and weeks. Your manager and the HR team are here to support you. Do not hesitate to ask questions at any point during your onboarding."),
        ("Before Your First Day", "You will receive an email with instructions to complete your pre-employment paperwork online. This includes your employment contract, tax forms, and bank details for payroll. You will also receive information about your first day including start time, location, and who to ask for upon arrival."),
        ("First Day", "On your first day, you will meet your manager, receive a tour of the office, and be introduced to your team. Your IT equipment will be ready for you. You will set up your accounts including email, Slack, and relevant software systems. The IT helpdesk is available at it-help@company.com for technical issues."),
        ("First Week", "During your first week, you will complete mandatory compliance training covering data protection, information security, and the code of conduct. These must be completed within 5 business days of starting. Your manager will schedule one-on-one meetings with key colleagues you will work with."),
        ("First Month", "In your first month, you will complete a full onboarding checklist provided by HR. You will have a 30-day check-in with your manager to discuss how onboarding is progressing. You will be assigned a buddy from your team who can answer informal questions."),
        ("HR Contacts", "For general HR queries contact hr@company.com. For payroll questions contact payroll@company.com. For benefits enrollment contact benefits@company.com. The HR team is available Monday to Friday 9am to 5pm. An HR portal is available at hr.company.com for self-service requests."),
        ("Probationary Period", "All new employees serve a 90-day probationary period. During this time, either party may terminate employment with 1 week's notice. At the end of the probationary period, your manager will conduct a review. Successful completion confirms your employment and activates your full benefits."),
    ]),
    ("docx", "data_privacy_security.docx", "Data Privacy and Security Policy", [
        ("Purpose", "This policy establishes requirements for protecting company data, client data, and personal data processed by the company. All employees must comply with this policy. Non-compliance may result in disciplinary action and may constitute a breach of applicable data protection law."),
        ("Password Requirements", "All passwords must be at least 12 characters long and include a mix of uppercase, lowercase, numbers, and symbols. Passwords must be changed every 90 days. Employees must not reuse the last 10 passwords. Passwords must never be shared with colleagues, written down, or stored in plain text."),
        ("Multi-Factor Authentication", "Multi-factor authentication (MFA) is mandatory for all company systems including email, VPN, cloud storage, and any system containing personal or confidential data. MFA must not be disabled. Loss of an MFA device must be reported to IT immediately."),
        ("Data Classification", "Company data is classified into four levels: Public (can be shared externally), Internal (for employees only), Confidential (restricted to those with need to know), and Restricted (highly sensitive, requires explicit authorization). Employee personal data and client data are classified as Confidential at minimum."),
        ("Incident Reporting", "Any suspected or confirmed security incident including data breaches, phishing attempts, malware, or loss of a device must be reported to the IT security team at security@company.com within 1 hour of discovery. Do not attempt to investigate or contain an incident without IT guidance. Under GDPR, breaches must be assessed within 72 hours."),
        ("Personal Devices", "Personal devices must not be used to access, store, or process company data or client data without explicit authorization from IT. Employees approved for BYOD must install the company's mobile device management software. Personal data from company systems must not be synced to personal cloud storage."),
        ("Data Retention and Deletion", "Data must only be retained for as long as necessary for its stated purpose and in accordance with the company's retention schedule. When data is no longer needed, it must be securely deleted using approved methods. Physical documents containing confidential data must be shredded, not placed in general waste."),
    ]),
    ("pdf", "parental_leave.pdf", "Parental Leave Policy", [
        ("Policy Statement", "The company is committed to supporting employees through significant life events including the birth or adoption of a child. This policy provides generous parental leave benefits above statutory minimums to help employees balance work and family responsibilities."),
        ("Primary Caregiver Leave", "The primary caregiver is entitled to 16 weeks of fully paid parental leave. This may be taken as a continuous block or in two separate periods within the first 12 months after the birth or adoption. Primary caregiver leave is available regardless of gender."),
        ("Secondary Caregiver Leave", "The secondary caregiver is entitled to 4 weeks of fully paid parental leave. This leave must be taken within 6 months of the birth or adoption. The 4 weeks may be taken consecutively or split into two separate 2-week periods."),
        ("Adoption Leave", "Employees who adopt a child are entitled to the same leave provisions as biological parents. Primary adopters receive 16 weeks paid leave and secondary adopters receive 4 weeks paid leave. Leave commences from the date the child is placed with the family."),
        ("Eligibility", "To be eligible for paid parental leave, employees must have completed at least 6 months of continuous employment at the time the leave is due to commence. Employees who do not meet the eligibility criteria may still be entitled to statutory parental leave."),
        ("Notice Requirements", "Employees should give at least 8 weeks notice of their intention to take parental leave where possible. Notice should be provided in writing to HR and the employee's manager. The expected start date of leave, type of leave, and anticipated return date must all be included."),
        ("Return to Work", "Employees returning from parental leave are entitled to return to the same role or, where that is not reasonably practicable, an equivalent role with the same terms and conditions. Employees should give 4 weeks notice of their intended return date. Flexible return arrangements may be requested and will be considered on a case-by-case basis."),
    ]),
    ("docx", "travel_policy.docx", "Business Travel Policy", [
        ("Purpose", "This policy governs all business travel by employees. All travel must be necessary for legitimate business purposes and must be approved in advance by the employee's manager. Employees are expected to exercise good judgment and spend company funds as they would their own."),
        ("Booking Process", "All business travel must be booked through the company's approved travel portal at travel.company.com. Booking outside the portal requires written pre-approval from a director. The travel portal provides access to negotiated rates and ensures compliance with this policy."),
        ("Air Travel", "Economy class is the standard for all flights. Business class is permitted for flights of 6 hours or more, or where the employee must be productive or client-facing immediately upon arrival. Upgrades using personal points or miles are permitted but the company will not reimburse upgrade fees."),
        ("Ground Transportation", "Employees should use the most cost-effective ground transportation. Taxis and ride-shares are reimbursable when public transport is unavailable or impractical. Car rentals should be economy class. Personal vehicle use is reimbursed at the approved mileage rate of EUR 0.30 per kilometer."),
        ("Per Diem", "The daily per diem for meals and incidentals is EUR 75 for domestic travel and EUR 100 for international travel. Actual meal receipts are not required when claiming per diem. If the company provides meals during the trip, the per diem is reduced proportionally."),
        ("Travel Insurance", "All employees traveling on company business are automatically covered by the company's corporate travel insurance. This covers medical emergencies, trip cancellation, lost luggage, and personal liability. Employees should carry their insurance card at all times during travel."),
        ("Expense Submission", "Travel expenses must be submitted within 5 business days of returning from the trip. Receipts for accommodation, flights, and any expense over EUR 25 must be attached. Late submissions may be declined. Intentional misrepresentation of expenses constitutes fraud and will result in immediate termination."),
    ]),
]


if __name__ == "__main__":
    for doc_type, filename, title, sections in DOCS:
        if doc_type == "pdf":
            make_pdf(filename, title, sections)
        else:
            make_docx(filename, title, sections)
    print(f"\nGenerated {len(DOCS)} documents in {OUTPUT_DIR}")
